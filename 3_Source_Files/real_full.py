import os
import cv2
import numpy as np
import pytesseract
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from matplotlib import pyplot as plt
import comtypes.client


output_path = "Outputs/"
input_path = "Images/"


def write_image(img,image_name,output_path):
	cv2.imwrite(output_path + image_name, img)

def replace_text(img):
	height, width, depth = img.shape
	imgScale = 1000/width
	newX,newY = img.shape[1]*imgScale, img.shape[0]*imgScale
	newimg = cv2.resize(img,(int(newX),int(newY)))
	charac=pytesseract.image_to_data(newimg)
	charlist=list(charac.split("\n"))
	for i in charlist:
		n=list(i.split("\t"))
		if len(n)==12 and n[10].isalpha()==False:
			if (n[11].isalpha()):
				newn6=int(int(n[6])/imgScale)
				newn7=int(int(n[7])/imgScale)
				newn8=int(int(n[8])/imgScale)
				newn9=int(int(n[9])/imgScale)
				cv2.rectangle(img,(newn6, newn7),((int((int(n[6]) + int(n[8]))/imgScale)), int((int(n[7]) + int(n[9]))/imgScale)), (255, 255, 255), -1)
				cv2.putText(img,n[11],(newn6,newn7+7), cv2.FONT_HERSHEY_SIMPLEX, 0.25, 0)
	finimg = cv2.resize(img,(int(img.shape[1]),int(img.shape[0])))
	return finimg
def crop_image(img,**kwargs):
	x1, y1 = kwargs['start']
	x2, y2 = kwargs['end']
	crop_img = img[y1:y2, x1:x2]
	return crop_img

def create_canny_edge(img):
	gray_image = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
	edged = cv2.Canny(gray_image, 0, 100)
	return edged

def vectorize(img):
	gray_image = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
	th3 = cv2.adaptiveThreshold(gray_image,255,cv2.ADAPTIVE_THRESH_GAUSSIAN_C,cv2.THRESH_BINARY,11,2)
	return th3

def find_contours(img):
	im0=img.copy()
	blur = cv2.GaussianBlur(img, (9, 9), 0)
	edged=create_canny_edge(blur)
	dilated = cv2.dilate(edged, np.ones((15, 15)))
	contours, hierarchy = cv2.findContours(dilated, cv2.RETR_EXTERNAL,cv2.CHAIN_APPROX_SIMPLE)
	boxes=[]
	for i, contour in enumerate(contours):
		rect = cv2.boundingRect(contour)
		x, y, w, h = [r*1 for r in rect]
		croped_image =crop_image(im0,**{'start': (x, y), 'end': (x+w, y+h)})
		boxes.append(croped_image)
		cv2.rectangle(img, (x,y), ((x+w), (y+h)), (200, 100, 150), 2)
	write_image(img,"boxed_image.png",output_path)
	return boxes

def detect_text(img):
	img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
	kernel = np.ones((1, 1), np.uint8)
	img = cv2.dilate(img, kernel, iterations=1)
	img = cv2.erode(img, kernel, iterations=1)
	result = pytesseract.image_to_string(img)
	return result

def prepare_text(text):
	list1=list(text.split("\n"))
	fin=""
	if len(list1)==0:
		return fin
	for i in range(len(list1)):
		fin+=list1[i]
		if len(list1[i])<60:
			fin+="\n"
		else:
			fin+=" "
	return fin

def check_for_sketch(img):
	img0=img.copy()
	length=img.shape[0]
	img1 = cv2.cvtColor(img0, cv2.COLOR_BGR2GRAY)
	kernel = np.ones((1, 1), np.uint8)
	img1 = cv2.dilate(img1, kernel, iterations=1)
	img1 = cv2.erode(img1, kernel, iterations=1)
	th3 = cv2.adaptiveThreshold(img1,255,cv2.ADAPTIVE_THRESH_GAUSSIAN_C,cv2.THRESH_BINARY,11,2)
	charac=pytesseract.image_to_data(img1)
	charlist=list(charac.split("\n"))
	for i in charlist:
		n=list(i.split("\t"))
		if len(n)==12 and n[10].isalpha()==False:
			if (n[11]!=""):
				cv2.rectangle(th3, (int(n[6])-5, int(n[7])-5), (int(n[6]) + int(n[8])+5, int(n[7]) + int(n[9])+5), (255, 255, 255), -1)
	blur = cv2.GaussianBlur(th3, (9, 9), 0)
	edged = cv2.Canny(blur, 0, 100)
	eroded = cv2.erode(edged, kernel, iterations=1)
	contours, hierarchy = cv2.findContours(eroded, cv2.RETR_EXTERNAL,cv2.CHAIN_APPROX_SIMPLE)
	for i, contour in enumerate(contours):
		rect = cv2.boundingRect(contour)
		x, y, w, h = [r*1 for r in rect]
		if w>50 or h>50:
			return True
	return False

def create_doc(image_path,output_type,output_name):
	ind=image_path.rfind('/')
	file_location=image_path[:ind]+'/'
	document= Document()
	image = cv2.imread(image_path)
	cv2.imwrite('Outputs/input.png',image)
	segments=find_contours(image)
	for i in range(len(segments)):
		segment=segments.pop()
		write_image(segment,"contoured_image%s.png" % i,output_path + "Contoured/")
		text=detect_text(segment)
		if check_for_sketch(segment):
			canny=vectorize(segment)
			write_image(canny,"canny_image%s.png" % i,output_path+ "Canny_Images/")
			document.add_picture(output_path+"Canny_Images/canny_image%s.png" %i)
			last_paragraph = document.paragraphs[-1]
			last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
		else:
			prepared_text=prepare_text(text)
			if prepared_text !="":
				document.add_paragraph(prepared_text)

	document.save(os.getcwd()+'\\Outputs\\Docs\\'+"%s.docx" % output_name)
	if output_type == "PDF":
		wdFormatPDF = 17
		word = comtypes.client.CreateObject('Word.Application')
		doc = word.Documents.Open(os.getcwd()+'\\Outputs\\Docs\\'+"%s.docx" % output_name)
		doc.SaveAs(file_location+"%s.PDF" % output_name, FileFormat=wdFormatPDF)
		doc.Close()
		word.Quit()
	elif output_type == "WORD":
		document.save(file_location+"%s.docx" % output_name)
	return
